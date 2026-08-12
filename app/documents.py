"""
Extraction du texte des documents ajoutes a la base de connaissances d'un
client (PDF, Word .docx, texte brut), pour servir de contexte supplementaire
a l'IA en complement du champ libre Client.contenu_site.

Seul le texte extrait est conserve (voir models.DocumentConnaissance) : le
fichier d'origine n'est jamais stocke.
"""

import io

import docx
import pypdf

EXTENSIONS_ACCEPTEES = {"pdf", "docx", "txt"}


def _extraire_pdf(contenu_octets: bytes) -> str:
    try:
        lecteur = pypdf.PdfReader(io.BytesIO(contenu_octets))
    except Exception as erreur:
        raise ValueError(f"Impossible de lire ce PDF (fichier corrompu ou protege) : {erreur}") from erreur

    if lecteur.is_encrypted:
        raise ValueError("Ce PDF est protege par mot de passe : impossible d'en extraire le texte.")

    morceaux = [page.extract_text() or "" for page in lecteur.pages]
    texte = "\n".join(morceaux).strip()
    if not texte:
        raise ValueError(
            "Aucun texte trouve dans ce PDF (probablement un scan sans reconnaissance de texte / OCR)."
        )
    return texte


def _extraire_docx(contenu_octets: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(contenu_octets))
    except Exception as erreur:
        raise ValueError(f"Impossible de lire ce fichier Word (.docx) : {erreur}") from erreur

    morceaux = [paragraphe.text for paragraphe in document.paragraphs]
    for tableau in document.tables:
        for ligne in tableau.rows:
            for cellule in ligne.cells:
                morceaux.append(cellule.text)

    texte = "\n".join(m for m in morceaux if m.strip()).strip()
    if not texte:
        raise ValueError("Aucun texte trouve dans ce document Word.")
    return texte


def extraire_texte(nom_fichier: str, contenu_octets: bytes) -> str:
    """
    Renvoie le texte extrait d'un fichier PDF, Word (.docx) ou texte brut.
    Leve ValueError avec un message clair si le format n'est pas pris en
    charge ou si l'extraction echoue.
    """
    extension = nom_fichier.rsplit(".", 1)[-1].lower() if "." in nom_fichier else ""

    if extension == "pdf":
        return _extraire_pdf(contenu_octets)
    if extension == "docx":
        return _extraire_docx(contenu_octets)
    if extension == "txt":
        texte = contenu_octets.decode("utf-8", errors="replace").strip()
        if not texte:
            raise ValueError("Ce fichier texte est vide.")
        return texte
    if extension == "doc":
        raise ValueError(
            "Le format .doc (ancien Word) n'est pas pris en charge : "
            "reenregistrez le document au format .docx puis reessayez."
        )
    raise ValueError(
        f"Format de fichier non pris en charge (.{extension or '?'}) : PDF, Word (.docx) ou texte (.txt) uniquement."
    )
